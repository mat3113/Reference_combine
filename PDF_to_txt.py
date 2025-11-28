import streamlit as st
import os
import io
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from bs4 import BeautifulSoup
from datetime import datetime # datetime import 추가

# -------------------- I. 상태 관리 초기화 --------------------

if 'content_sections' not in st.session_state:
    # 본문 내용 저장: {'title', 'text', 'source_ref'}
    st.session_state.content_sections = []  
if 'references' not in st.session_state:
    # 출처 목록 저장: {'type', 'source', 'link'}
    st.session_state.references = []      
if 'ref_count' not in st.session_state:
    st.session_state.ref_count = 1

# -------------------- II. 핵심 기능 함수 --------------------

def extract_text_from_pdf(pdf_bytes):
    """PDF 바이트에서 텍스트를 추출합니다."""
    text = ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                # 페이지 구분을 명확히
                text += page_text + "\n\n--- 페이지 구분 ---\n\n" 
        return text.strip()
    except Exception as e:
        st.error(f"PDF 텍스트 추출 중 오류 발생: {e}")
        return ""

def get_url_title(url):
    """URL에서 페이지 제목을 추출합니다."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=5)
        response.raise_for_status() 
        soup = BeautifulSoup(response.content, 'html.parser')
        title = soup.title.string if soup.title else url
        return title.strip()
    except requests.exceptions.RequestException as e:
        return f"[링크 접근 실패] {url}"
    except Exception:
        return url

def add_reference(ref_type, source, link):
    """출처 목록에 새 항목을 추가하고 고유 참조 태그를 반환합니다."""
    # '출처를 직접 입력하세요' 태그는 중복 검사에서 제외
    is_manual_input = (link == "출처를 직접 입력하세요")
    
    # 중복 방지를 위해 이미 같은 링크가 있는지 확인 (수동 입력 제외)
    if not is_manual_input and any(ref['link'] == link for ref in st.session_state.references):
        st.warning(f"이미 추가된 출처입니다: {link}")
        return None 

    ref_tag = f"[{st.session_state.ref_count}]"
    
    st.session_state.references.append({
        'type': ref_type,
        'source': source,
        'link': link
    })
    st.session_state.ref_count += 1
    return ref_tag

def create_docx_report():
    """DOCX 보고서를 생성하고 io.BytesIO 객체로 반환합니다."""
    doc = Document()
    
    # 폰트 및 스타일 설정 (기본적으로 맑은 고딕 사용 가정)
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    # 1. 보고서 제목
    doc.add_heading("연구 자료 정리 보고서", level=0)

    # 2. 본문 섹션
    doc.add_heading("1. 자료 본문 정리", level=1)
    
    if st.session_state.content_sections:
        for section in st.session_state.content_sections:
            # 섹션 제목
            doc.add_heading(section['title'], level=2)
            
            # 본문 내용
            for paragraph in section['text'].split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            # 출처 태그 추가
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(f"(출처: {section['source_ref']})").italic = True
            doc.add_page_break()
    else:
        doc.add_paragraph("정리된 본문 내용이 없습니다.")

    # 3. 출처 목록
    doc.add_heading("2. 출처 (References)", level=1)
    
    if st.session_state.references:
        # 출처 목록을 정리된 순서대로 표시
        for i, ref in enumerate(st.session_state.references):
            ref_tag = f"[{i + 1}]"
            source_info = ref['source']
            link_info = ref['link']
            
            p = doc.add_paragraph()
            p.add_run(ref_tag).bold = True
            p.add_run(f" {source_info} ")
            p.add_run(f"({ref['type']} 자료)").italic = True
            
            # 링크 정보 처리
            if link_info == "출처를 직접 입력하세요":
                 p.add_run(f" - {link_info} (제목: {source_info})").italic = True
            elif len(link_info) > 80:
                 p.add_run(f"\n   링크/경로: {link_info}")
            else:
                 p.add_run(f" - 링크/경로: {link_info}")
    else:
        doc.add_paragraph("기록된 출처가 없습니다.")

    # 파일 저장 및 반환
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# -------------------- IV. Streamlit UI --------------------

st.title("📚 연구 자료 정리 및 출처 관리 도구")
st.markdown("---")

# -------------------- A. 자료 입력 사이드바 --------------------
with st.sidebar:
    st.header("자료 입력 📥")
    
    # 1. PDF 논문 텍스트 입력
    with st.expander("1. PDF 논문 텍스트 추출", expanded=True):
        pdf_file = st.file_uploader("PDF 파일 업로드", type=["pdf"])
        pdf_title = st.text_input("논문 제목 (미입력 시 파일명 사용)", key='pdf_title')
        
        if st.button("PDF 텍스트 추출 및 추가", key='btn_pdf'):
            if pdf_file:
                with st.spinner("PDF 텍스트 추출 중..."):
                    pdf_bytes = pdf_file.read()
                    extracted_text = extract_text_from_pdf(pdf_bytes)
                
                if extracted_text:
                    title = pdf_title if pdf_title else os.path.splitext(pdf_file.name)[0]
                    ref_tag = add_reference('PDF 논문', title, pdf_file.name)
                    
                    if ref_tag:
                        st.session_state.content_sections.append({
                            'title': title,
                            'text': extracted_text,
                            'source_ref': ref_tag
                        })
                        st.success(f"PDF({title}) 내용이 본문에 추가되었습니다. 출처: {ref_tag}")
                    # else: add_reference에서 이미 warning 출력
                else:
                    st.error("추출된 텍스트가 없거나 오류가 발생했습니다.")
            else:
                st.error("PDF 파일을 먼저 업로드해주세요.")

    # 2. 영상 링크/HTML 입력
    with st.expander("2. 영상/HTML 출처 링크 추가"):
        url_link = st.text_input("URL 입력 (영상, 기사, HTML)", key='url_link')
        url_type = st.selectbox("자료 유형", ['영상 링크', 'HTML 기사', '기타 웹문서'], key='url_type')
        
        if st.button("URL 출처 목록에 추가", key='btn_url'):
            if url_link:
                with st.spinner("URL 제목을 가져오는 중..."):
                    title = get_url_title(url_link)
                
                ref_tag = add_reference(url_type, title, url_link)
                
                if ref_tag:
                    st.success(f"출처가 추가되었습니다: {title} {ref_tag}")
                # else: add_reference에서 이미 warning 출력
            else:
                st.error("URL을 입력해주세요.")

    # 3. 이미지 및 수동 입력
    with st.expander("3. 이미지 및 수동 출처/내용 입력"):
        manual_content = st.text_area("본문 내용 (직접 입력)", height=150, key='manual_content')
        manual_source = st.text_input("출처 제목 (예: Figure 1. 인공지능 이미지)", key='manual_source')
        # 출처 판단이 힘들면 '출처를 직접 입력하세요' 태그를 사용
        manual_link_input = st.text_input("출처 링크 (없거나 판단 불가 시 공란으로 두세요)", key='manual_link_input')
        
        if st.button("수동 내용 및 출처 추가", key='btn_manual'):
            if manual_content and manual_source:
                link = manual_link_input if manual_link_input else "출처를 직접 입력하세요"
                
                ref_tag = add_reference('이미지/수동 자료', manual_source, link)
                
                if ref_tag:
                    st.session_state.content_sections.append({
                        'title': manual_source,
                        'text': manual_content,
                        'source_ref': ref_tag
                    })
                    st.success(f"수동 내용이 본문에 추가되었습니다. 출처: {ref_tag}")
                # else: add_reference에서 이미 warning 출력
            else:
                st.error("내용과 출처 제목을 모두 입력해야 합니다.")


# -------------------- B. 메인 화면: 현재 상태 및 DOCX 출력 --------------------

st.header("현재 정리 상태")
st.markdown("---")

# 1. 출처 목록 (References)
st.subheader("1. 출처 목록")
if st.session_state.references:
    ref_df = []
    for i, ref in enumerate(st.session_state.references):
        ref_df.append({
            'Tag': f"[{i + 1}]",
            '유형': ref['type'],
            '제목/내용': ref['source'],
            '링크/경로': ref['link']
        })
    # Streamlit DataFrame으로 시각화하여 보여줌
    st.dataframe(ref_df, use_container_width=True, hide_index=True)
else:
    st.info("아직 추가된 출처가 없습니다. 사이드바를 이용해 자료를 추가해주세요.")

# 2. 본문 내용 (Content Sections)
st.subheader("2. 본문 정리 섹션")
if st.session_state.content_sections:
    section_titles = [f"{s['title']} (출처: {s['source_ref']})" for s in st.session_state.content_sections]
    st.success(f"총 {len(section_titles)}개의 섹션이 DOCX에 정리됩니다.")
    with st.expander("추가된 섹션 제목 및 출처 확인"):
        st.write(section_titles)
    
    # 초기화 버튼
    if st.button("정리된 자료 초기화"):
        st.session_state.content_sections = []
        st.session_state.references = []
        st.session_state.ref_count = 1
        st.rerun() # 상태를 초기화하고 앱을 새로고침
else:
    st.info("아직 정리된 본문 내용이 없습니다.")


# 3. DOCX 다운로드
st.header("3. DOCX 보고서 생성 및 다운로드")
if st.session_state.content_sections:
    # DOCX 생성 함수 호출 (버튼 클릭 시에만 생성)
    docx_io = create_docx_report()
    
    st.download_button(
        label="최종 보고서 다운로드 (.docx)",
        data=docx_io,
        file_name=f"Research_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key='download_docx'
    )
    st.success("다운로드 버튼을 클릭하여 보고서를 저장하세요.")
else:
    st.warning("DOCX를 생성하려면 하나 이상의 자료를 추가해야 합니다.")
